"""Fill public/case-template.docx with Tulp case data."""
import json
import sys
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

TEMPLATE = r'C:\Users\ian_s\repos\creates\Sales Navigator\public\case-template.docx'
JSON_PATH = r'C:\Users\ian_s\repos\creates\Sales Navigator\tmp\tulp_case.json'
OUTPUT = r'C:\Users\ian_s\repos\creates\Sales Navigator\tmp\case-tulp-group.docx'

with open(JSON_PATH, encoding='utf-8') as f:
    data = json.load(f)

doc = Document(TEMPLATE)

def set_cell_text(cell, text):
    """Replace cell text while keeping the first paragraph's formatting."""
    # Keep first paragraph, clear others
    paragraphs = cell.paragraphs
    first = paragraphs[0]
    # remove extra paragraphs
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)
    # clear runs in first paragraph (keep first run formatting)
    runs = first.runs
    if runs:
        first_run = runs[0]
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
        first_run.text = text
    else:
        first.add_run(text)

# Table 0: bedrijfsnaam + korte omschrijving
t0 = doc.tables[0]
set_cell_text(t0.rows[0].cells[1], data['bedrijfsnaam'])
set_cell_text(t0.rows[1].cells[1], data['korte_omschrijving'])

# Tables 1..6: single-cell content
single_fields = ['situatie', 'doel', 'oplossing', 'resultaat', 'keywords', 'business_impact']
for i, key in enumerate(single_fields, start=1):
    tbl = doc.tables[i]
    set_cell_text(tbl.rows[0].cells[0], data[key])

# Tables 7,8,9: mapping (doelen/behoeften/diensten)
def apply_mapping(tbl, mapping):
    for row in tbl.rows[1:]:
        label_cell = row.cells[0]
        toel_cell = row.cells[1]
        original = label_cell.text
        # Find matching key by substring (option label appears after ☐)
        matched_key = None
        for k in mapping:
            if k in original:
                matched_key = k
                break
        if matched_key is None:
            continue
        toel = mapping[matched_key]
        if toel.strip():
            # Replace ☐ with ☑
            new_label = original.replace('☐', '☑')
            set_cell_text(label_cell, new_label)
            set_cell_text(toel_cell, toel)

apply_mapping(doc.tables[7], data['mapping']['doelen'])
apply_mapping(doc.tables[8], data['mapping']['behoeften'])
apply_mapping(doc.tables[9], data['mapping']['diensten'])

doc.save(OUTPUT)
print(f'OK -> {OUTPUT}')
